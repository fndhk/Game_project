from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "DarkUs_GameDesignDocument_v0.1.docx"


FONT = "Malgun Gothic"
TITLE_COLOR = RGBColor(11, 37, 69)
H1_COLOR = RGBColor(46, 116, 181)
H3_COLOR = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 110)
TABLE_HEADER = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "B9C4D0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_run(run, size=None, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, 28, True, TITLE_COLOR)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(subtitle)
    style_run(run, 12, False, MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    if level == 1:
        style_run(run, 16, True, H1_COLOR)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        style_run(run, 13, True, H1_COLOR)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        style_run(run, 12, True, H3_COLOR)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    style_run(run, 11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    style_run(run, 11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    style_run(run, 11)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, "D6DEE8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    style_run(run, 11, True, TITLE_COLOR)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(body)
    style_run(run, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_kv_table(doc, rows, widths=(1.6, 4.9)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_column_widths(table, widths)
    set_table_borders(table)
    for idx, (key, value) in enumerate(rows):
        c0, c1 = table.rows[idx].cells
        set_cell_shading(c0, TABLE_HEADER)
        set_cell_text(c0, key, bold=True, color=TITLE_COLOR)
        set_cell_text(c1, value)
        for cell in table.rows[idx].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_column_widths(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_text(cell, header, bold=True, color=TITLE_COLOR)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build():
    doc = Document()
    configure_document(doc)

    add_title(
        doc,
        "dark us v-1 게임 기획서",
        "Game Design Document 초안 v0.1 | 작성 기준: 현재 Unity 프로젝트 구조 및 스크립트 분석",
    )
    add_callout(
        doc,
        "기획 핵심",
        "어두운 연구소에서 직접 시야 대신 라이다 스캔 점으로 공간을 파악하고, 시민은 목표 컴퓨터를 복구해 탈출하며, 킬러는 제한된 킬타임 동안 시민을 제거하는 비대칭 멀티플레이 공포 게임.",
    )
    add_kv_table(
        doc,
        [
            ("문서 목적", "현재 구현 상태를 바탕으로 게임 개요, 시스템, 캐릭터/시나리오, 레벨, 그래픽/사운드, 테스트, 출시 계획을 한 문서로 정리한다."),
            ("적용 범위", "기획 초안, 팀 공유용, 포트폴리오/발표 자료의 원본 문서"),
            ("확정 필요", "정식 게임명, 최종 출시 플랫폼, 가격, 연령 등급, 최종 아트 콘셉트"),
        ],
    )

    add_heading(doc, "목차", 1)
    for item in [
        "1. 게임 개요",
        "2. 게임 시스템 디자인",
        "3. 캐릭터 및 시나리오",
        "4. 레벨 디자인 설계",
        "5. 게임 그래픽 아트 디자인",
        "6. 게임 사운드 디자인",
        "7. 게임 플레이 테스트 계획",
        "8. 게임 출시 및 마케팅 계획",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1. 게임 개요", 1)
    add_kv_table(
        doc,
        [
            ("게임 타이틀", "dark us v-1 (가제)"),
            ("게임 장르", "1인칭 비대칭 멀티플레이 공포 / 생존 / 소셜 디덕션"),
            ("게임 플랫폼", "우선 대상: PC Windows. 엔진: Unity. 네트워크: Photon PUN 기반 온라인 멀티플레이."),
            ("게임 설명", "플레이어는 빛이 차단된 지하 연구소에 투입된다. 시민은 라이다 스캐너로 보이지 않는 공간을 점으로 드러내며 목표 컴퓨터를 복구하고 탈출구를 열어야 한다. 킬러는 같은 플레이어처럼 섞여 있다가 정해진 킬타임에 시민을 제거한다."),
            ("게임 목표", "시민: 목표 컴퓨터 또는 Access Core 목표를 완료하고 탈출한다. 킬러: 시민이 탈출하기 전에 모두 제거하거나 제한 시간을 버틴다."),
        ],
    )
    add_heading(doc, "핵심 경험", 2)
    for item in [
        "완전한 시야가 아니라 스캔 점으로만 공간과 대상의 윤곽을 읽는 긴장감",
        "연구소 목표 수행, 동선 판단, 탈출 타이밍을 둘러싼 협동 플레이",
        "킬러가 시민 사이에 숨어 있다가 제한된 시간 창에서만 직접적인 살해 행동을 할 수 있는 압박감",
        "무작위 연구소 생성과 Photon 동기화를 통해 매 라운드 다른 동선과 의심 구도를 제공",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 게임 시스템 디자인", 1)
    add_heading(doc, "게임 진행 방식", 2)
    for item in [
        "로비에서 방을 만들거나 공개 방 목록을 통해 입장한다.",
        "게임 시작 시 Photon 방 속성에 저장된 임포스터 Actor 번호를 기준으로 킬러 1명과 시민 다수를 배정한다.",
        "라운드 시작 후 20분 제한 시간이 흐른다.",
        "시민은 목표 컴퓨터를 찾아 복구한다. 설정에 따라 Access Core 수집 및 설치 방식도 지원한다.",
        "필요 목표 수를 달성하면 비상 탈출문이 잠금 해제되고, 시민이 탈출하면 시민 승리 조건을 평가한다.",
        "킬러는 킬타임 동안 Q 키로 근접 즉사 공격을 실행할 수 있다.",
        "시간이 만료되거나 살아서 탈출 가능한 시민이 없으면 킬러 승리로 처리한다.",
    ]:
        add_number(doc, item)

    add_heading(doc, "게임 룰", 2)
    add_matrix(
        doc,
        ["항목", "현재 설계값", "기획 의도"],
        [
            ("라운드 시간", "20분", "탐색과 의심, 목표 수행을 모두 경험할 수 있는 중간 길이의 세션"),
            ("킬타임", "18~15분, 13~10분, 8~5분, 3~0분 남은 구간", "라운드 내내 공격 가능하지 않게 하여 은신과 타이밍 싸움을 강화"),
            ("시민 승리", "필요 시민 탈출 수 달성", "목표 수행 이후 탈출 행동까지 완료해야 승리"),
            ("킬러 승리", "시간 만료 또는 시민 전원 사망/탈출 불가", "소극적 방해와 직접 처치 모두 승리 루트로 인정"),
            ("스캔 입력", "스캔 키를 누르는 동안 쿨타임마다 펄스 발생", "짧은 순간에만 공간 정보가 드러나는 제한된 인지 경험"),
        ],
        [1.35, 2.15, 3.0],
    )
    add_heading(doc, "게임 인터페이스 UI/UX", 2)
    for item in [
        "HUD는 라이다 장비 콘셉트의 얇은 선, 어두운 패널, 청록/백색 계열 정보를 사용한다.",
        "중앙에는 스캔 쿨타임과 조준 보조 요소가 표시된다.",
        "상단 중앙에는 라운드 타이머가 표시되며 킬타임에는 붉은 경고 색상과 메시지로 전환된다.",
        "상단 좌측에는 dot memory, 하단 좌측에는 생명/스태미나, 하단 중앙에는 아이템 슬롯을 배치한다.",
        "우측 상단에는 목표 진행도: Find Target Computers 0/4 같은 형태로 표시한다.",
        "음성 채팅 상태는 MIC OPEN/MIC MUTED로 짧게 표시해 협동과 의심 상황을 보조한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 캐릭터 및 시나리오", 1)
    add_heading(doc, "주요 캐릭터 소개", 2)
    add_matrix(
        doc,
        ["역할", "플레이 목적", "주요 행동"],
        [
            ("시민", "연구소의 목표를 복구하고 탈출한다.", "스캔, 탐색, 컴퓨터 복구, 아이템 사용, 탈출문 상호작용"),
            ("킬러", "시민의 목표 진행을 늦추고 킬타임에 시민을 제거한다.", "시민처럼 행동, 추적, 킬타임 공격, 혼란 유도"),
            ("시스템/연구소", "공간 자체가 위협이 되는 무대 역할을 한다.", "시야 차단, 무작위 동선, 목표 위치 분산, 탈출구 제한"),
        ],
        [1.2, 2.7, 2.6],
    )
    add_heading(doc, "스토리 개요", 2)
    add_body(doc, "플레이어들은 통신이 끊긴 지하 연구소의 복구 팀으로 투입된다. 시설 내부는 정전과 알 수 없는 실험의 영향으로 일반적인 시야 확보가 불가능하다. 팀은 라이다 스캐너로 표면을 드러내며 목표 컴퓨터를 복구해야 하지만, 복구 팀 안에는 연구소 사고의 원인과 연결된 변이체 또는 배신자가 숨어 있다.")
    add_body(doc, "라운드의 서사는 매번 같다. 공간을 읽고, 단서를 찾고, 목표를 고치고, 탈출문까지 이동한다. 그러나 누가 킬러인지, 목표와 탈출구가 어디에 있는지, 누가 마지막까지 살아남을지는 매 라운드 달라진다.")
    add_heading(doc, "배경 설정", 2)
    for item in [
        "장소: 지하 연구소, 실험실, 복도, 보안 구역, 비상 탈출구",
        "상태: 조명 대부분이 꺼져 있으며, 실제 맵 렌더러는 숨겨지고 충돌체와 스캔 점으로 공간을 파악한다.",
        "위협: 킬러의 정체는 시작 시 숨겨지고, 킬타임 경고가 뜨는 순간 플레이어 간 신뢰가 급격히 흔들린다.",
        "목표: 복구해야 하는 목표 컴퓨터는 스캔 표면 색상과 상호작용 안내로 구분한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 레벨 디자인 설계", 1)
    add_heading(doc, "스테이지 레벨 목록", 2)
    add_kv_table(
        doc,
        [
            ("현재 메인 스테이지", "Procedural Underground Laboratory"),
            ("기본 구성", "방, 복도, 계단/수직 연결, 시작방, 목표 컴퓨터, 아이템 스폰 포인트, 탈출문"),
            ("기본 생성값", "RoomCount 30, PlayerCount 4, MapSize 80x80, FloorHeight 4"),
            ("동기화", "Photon 방 seed를 사용해 클라이언트 간 동일한 맵을 생성"),
        ],
    )
    add_heading(doc, "각 레벨의 목적과 특징", 2)
    add_matrix(
        doc,
        ["공간 요소", "역할", "기획 포인트"],
        [
            ("시작방", "플레이어 초기 위치", "플레이어끼리 같은 방에서 시작하지 않도록 설정 가능"),
            ("일반 방", "목표/아이템/은신 공간", "스캔으로 내부 구조를 파악하고 위험을 감수해 진입"),
            ("복도", "이동과 추적", "시야가 좁고 마주침이 자주 발생하는 긴장 구간"),
            ("계단/수직 연결", "층 이동", "도주와 우회 루트를 제공"),
            ("탈출구", "최종 목표", "시작방에서 일정 거리 이상 떨어진 곳에 배치해 이동 리스크를 보장"),
        ],
        [1.35, 2.0, 3.15],
    )

    add_heading(doc, "5. 게임 그래픽 아트 디자인", 1)
    add_heading(doc, "캐릭터 디자인", 2)
    for item in [
        "시민은 장비를 착용한 연구소 복구 인원으로 설계한다.",
        "킬러는 평상시 시민과 구분되지 않되, 공격/스캔/사망 연출에서만 위협성이 드러나게 한다.",
        "멀티플레이 식별을 위해 플레이어별 스캔 색상 그룹을 사용한다.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "배경 디자인", 2)
    for item in [
        "연구소 본체는 인게임에서 직접 보이지 않거나 매우 제한적으로만 노출된다.",
        "공간의 실루엣은 라이다 점, 금속/유리/벽/바닥 표면 색상, 목표 오브젝트 색상으로 전달한다.",
        "검은 배경 위에 점과 HUD가 떠오르는 미니멀 공포 톤을 유지한다.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "효과 디자인", 2)
    add_matrix(
        doc,
        ["효과", "표현 방식", "목적"],
        [
            ("스캔 펄스", "화면 넓은 범위에 점이 파동처럼 생성", "보이지 않는 지형과 오브젝트를 순간적으로 인지"),
            ("표면 색상", "바닥, 벽, 금속, 유리, 목표, 플레이어, 아이템 별 색상 그룹", "정보 전달과 플레이 판단 보조"),
            ("킬타임 경고", "타이머/문구 붉은색 전환", "위험 구간을 즉시 인지"),
            ("탈출문", "잠금 해제 후 문 이동 및 충돌 해제", "목표 완료의 명확한 보상 피드백"),
        ],
        [1.3, 2.6, 2.6],
    )

    add_heading(doc, "6. 게임 사운드 디자인", 1)
    add_heading(doc, "배경음악 BGM", 2)
    for item in [
        "기본 BGM은 저주파 드론, 기계음, 멀리서 들리는 환풍/전력 노이즈 중심으로 구성한다.",
        "킬타임에는 리듬이 강한 음악보다 심박, 경고음, 저역 압박감을 높여 긴장도를 상승시킨다.",
        "로비는 인게임보다 덜 위협적인 전자음으로 분리해 준비 단계와 플레이 단계를 명확히 구분한다.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "효과음 SFX", 2)
    for item in [
        "스캔 펄스음은 로컬 재생과 네트워크 위치 재생을 모두 사용해 다른 플레이어에게 위치 정보를 남긴다.",
        "발소리는 바닥 재질과 거리에 따라 작은 단서가 되도록 설계한다.",
        "컴퓨터 복구, 아이템 획득, 탈출문 잠금 해제, 문 개방에는 짧고 명확한 확인음을 배치한다.",
        "킬러 공격은 짧은 선딜 후 즉시성이 느껴지는 저음 타격음으로 처리한다.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "목소리 Voice", 2)
    for item in [
        "플레이어 간 음성 채팅을 기본 커뮤니케이션 수단으로 둔다.",
        "마이크 상태 UI를 제공해 음소거 여부를 즉시 알 수 있게 한다.",
        "거리 기반 음성 또는 팀 전체 음성 여부는 테스트를 통해 최종 결정한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. 게임 플레이 테스트 계획", 1)
    add_heading(doc, "플레이 테스트 계획", 2)
    add_matrix(
        doc,
        ["테스트 항목", "확인 내용", "성공 기준"],
        [
            ("기본 루프", "로비 입장, 역할 배정, 목표 복구, 탈출/사망, 결과 후 로비 복귀", "한 라운드를 중단 없이 완료"),
            ("스캔 가독성", "점 밀도, 색상 구분, 쿨타임, 성능", "플레이어가 길/목표/위협을 읽을 수 있고 프레임 저하가 크지 않음"),
            ("킬타임 밸런스", "킬러 공격 가능 구간과 시민 생존률", "시민과 킬러 모두 승산을 체감"),
            ("네트워크 동기화", "맵 seed, 역할, 사망, 탈출, 컴퓨터 복구 이벤트", "클라이언트 간 상태 불일치 없음"),
            ("UI/UX", "목표, 타이머, 스캔, 아이템, 음성 상태", "초보 플레이어가 1라운드 안에 목표를 이해"),
        ],
        [1.35, 3.0, 2.15],
    )
    add_heading(doc, "게임 디버깅 계획", 2)
    for item in [
        "Unity 콘솔 로그 기준: 맵 생성 실패, Photon 역할 미할당, GameOver 중복 호출, 목표 컴퓨터 인덱스 누락을 우선 추적한다.",
        "멀티플레이는 최소 2클라이언트, 권장 4클라이언트로 반복 검증한다.",
        "라이다 점 누적량과 GPU 인스턴싱 렌더러 메모리를 장시간 플레이 기준으로 관찰한다.",
        "탈출문 잠금 해제, 사망 처리, 로비 복귀는 모든 클라이언트에서 같은 결과가 나오는지 확인한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. 게임 출시 및 마케팅 계획", 1)
    add_kv_table(
        doc,
        [
            ("출시 일정", "미정. 권장 초안: 2026년 하반기 프로토타입 공개, 이후 플레이 테스트 결과에 따라 얼리 액세스 여부 결정."),
            ("게임 가격", "미정. 소규모 온라인 공포 게임 기준 무료 데모 + 저가 유료판 또는 무료 공개 후 후원/포트폴리오 활용을 검토."),
            ("게임 마케팅 전략", "라이다 스캔으로만 보이는 연구소, 킬타임 배신 구조, 4인 협동/의심 플레이를 짧은 영상으로 강조한다."),
        ],
    )
    add_heading(doc, "마케팅 메시지 초안", 2)
    for item in [
        "빛이 아니라 점으로 길을 찾는다.",
        "복구 팀 안에 킬러가 있다.",
        "스캔할수록 위치가 드러나고, 숨을수록 길을 잃는다.",
        "20분 안에 연구소를 고치고 탈출하라.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "부록: 현재 구현 기반 메모", 1)
    add_kv_table(
        doc,
        [
            ("확인한 주요 스크립트", "GameLoopManager, RoleAssignmentManager, RoundTimer, LidarSpotScanner, LabObjectiveManager, EmergencyExitDoor, KillerAttack, LaboratoryGenerator, PlayerHUDController"),
            ("현재 강점", "핵심 루프와 스캔 시각화, Photon 동기화, 절차적 연구소 생성, HUD가 이미 구현 방향을 갖고 있음"),
            ("다음 보완 과제", "정식 아트 스타일 확정, 사운드 리소스 확정, 튜토리얼/온보딩, 밸런스 수치 조정, 매치메이킹 안정화"),
        ],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
